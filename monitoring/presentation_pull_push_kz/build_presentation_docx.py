# -*- coding: utf-8 -*-
"""Генерирует Word-документ со слайдами, текстом и указанием спикера."""
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Установите: pip install python-docx")
    raise

OUT = "presentation_pull_push_kz.docx"


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def page_break(doc):
    doc.add_page_break()


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    slides = [
        {
            "n": 1,
            "title": "Титул",
            "speakers": "Спикер 1 + Спикер 2 (бірге)",
            "slide_text": [
                "Тақырып: Pull және Push модельдері: Graphite, Nagios, Zabbix және Prometheus мысалындағы мониторинг",
                "Қосымша: Орындаушылар: [толтырыңыз] | Жобалар: Fit AI (Windows), Volt Store / стендтер (Ubuntu)",
            ],
            "script": (
                "Құрметті комиссия мүшелері! Біздің баяндамамыз мониторинг жүйелерінде метрикаларды жинаудың екі негізгі "
                "тәсілін — Pull және Push модельдерін салыстыруға арналған. Біз теорияны ғана емес, Graphite, Nagios, "
                "Zabbix және Prometheus құралдарының нақты мысалдарын көрсетеміз. Серіктесім Ubuntu ортасында Graphite "
                "пен Nagios бойынша практикалық жұмысты орындады, ал мен Windows ортасында Zabbix стендін Docker Compose "
                "арқылы көтердім. Prometheus бойынша Pull-логикасын біздің бұрынғы жобамыздағы мониторингпен байланыстырамыз."
            ),
            "screens": "Скрин: міндетті емес. Қалағанда: Docker / Prometheus / Grafana / Zabbix логотиптерінің коллажы.",
        },
        {
            "n": 2,
            "title": "Мақсат, міндеттер, күтілетін нәтиже",
            "speakers": "Негізінен Спикер 1 (негізгі мәтін)",
            "slide_text": [
                "Мақсат: Pull және Push идеяларын нақты құралдармен салыстыру",
                "Міндеттер:",
                "  • Push моделін Graphite мысалында көрсету (Спикер 2)",
                "  • Белсенді тексерулерді Nagios мысалында көрсету (Спикер 2)",
                "  • Zabbixты Docker арқылы орнату және UI-де растау (Спикер 1)",
                "  • Prometheus Pull scrape логикасын Fit AI мониторингпен байланыстыру (Спикер 1)",
                "Нәтиже: әр жүйенің қолдану аясы мен модельдік ерекшелігі түсінікті болады",
            ],
            "script": (
                "Біздің мақсатымыз — «метрикаларды қалай жинаймыз» деген сұраққа тек анықтама бермей, нақты сценарийлер "
                "арқылы жауап беру. Біріншіден, деректерді қолданбаның өзі жіберетін Push жағына Graphite арқылы "
                "серіктесім тоқталады. Екіншіден, мониторинг серверінің уақыт бойынша тексеретін белсенді схемасына "
                "Nagios мысалын қосамыз. Үшіншіден, кәсіби мониторинг платформасы ретінде Zabbix стендін мен көрсетемін. "
                "Сонында Prometheus бойынша Pull scrape принципін біздің Fit AI жобасындағы бақылаумен ұштастырамыз."
            ),
            "screens": "Скрин: қажет емес немесе жалпы блок-схема (опционал).",
        },
        {
            "n": 3,
            "title": "Pull vs Push: анықтамалар",
            "speakers": "Спикер 2",
            "slide_text": [
                "Push: метрикаларды дереккөз өзі жібереді (мысалы Carbon портына жазу)",
                "Pull: мониторинг сервері мақсатқа барып, метрикаларды өзі алады (Prometheus scrape)",
                "Nagios: белсенді тексерулер — уақыт бойынша опрос; бұл push емес",
                "Zabbix: гибридті схема — агент + жоспарланған тексерулер",
            ],
            "script": (
                "Push моделінде бастаманы жіберуші жасайды: Graphite экожүйесінде жиі Carbon арқылы plaintext хабарламасымен "
                "жүзегеседі. Pull моделінде бастаманы мониторинг сервері алады: Prometheus мақсаттар тізімін біледі және "
                "/metrics сияқты нүктелерге scrape жасайды. Nagios үшін негізгі идея — белсенді тексерулер: сервер хост пен "
                "сервистерді күтпей, сценарий бойынша сұрайды. Zabbixта екі әлемнің элементтері бір платформада кездеседі."
            ),
            "screens": "Скрин: кіші диаграмма (Push қолданба→Carbon; Pull Prometheus→/metrics).",
        },
        {
            "n": 4,
            "title": "Graphite (Ubuntu, Docker)",
            "speakers": "Спикер 2",
            "slide_text": [
                "Контейнер: graphiteapp/graphite-statsd",
                "Метрика: volt-store.cpu → Carbon, plaintext, порт 2003",
                "Graphite Web UI: порт 8080 — уақыт қатары көрінеді",
                "Қорытынды: Push моделі — деректерді жіберу басталады",
            ],
            "script": (
                "Ubuntu ортасында біз Graphite стекін Docker арқылы көтердік. Деректер Carbon арқылы қабылданады: "
                "метриканы plaintext форматында жібердік, ол уақыт белгісімен уақыт қатарына сақталады. Graphite Web "
                "интерфейсінде графикті көрдік. Бұл Push моделінің нақты демонстрациясы: бастаманы дерек жіберуші жасайды."
            ),
            "screens": "1) Graphite Web — график volt-store.cpu. 2) Терминал — жіберу командасы (nc/echo).",
        },
        {
            "n": 5,
            "title": "Nagios (Ubuntu, Docker)",
            "speakers": "Спикер 2",
            "slide_text": [
                "Контейнер: jasonrivers/nagios, веб: порт 8081",
                "Хост: volt-store (Graphite контейнеріне сілтеме)",
                "Service check: HTTP — OK, жауап 200",
                "Қорытынды: белсенді мониторинг — уақыт бойынша тексеру",
            ],
            "script": (
                "Nagios үшін контейнерлік шешім қолданылды. Хост ретінде Graphite контейнеріне сілтеме жасалды, HTTP "
                "тексеру таңдалды: сервис қолжетімді ме және дұрыс жауап келе ме. Нәтиже: хост UP, сервис OK, HTTP 200. "
                "Бұл классикалық белсенді тексеру: Nagios өзі сұраныс жасайды."
            ),
            "screens": "1) Nagios — хост UP, сервис OK. 2) Қосымша: service details.",
        },
        {
            "n": 6,
            "title": "Prometheus: Pull моделі (Fit AI)",
            "speakers": "Спикер 1",
            "slide_text": [
                "Prometheus мақсаттарды scrape жасайды",
                "Backend /metrics — қолданба метрикалары",
                "Салыстыру: Graphite Push ↔ Prometheus Pull",
            ],
            "script": (
                "Prometheus біздің Fit AI жобасында бұрыннан қолданылады: ол backend /metrics endpoint-іне уақыт аралығымен "
                "барып, метрикаларды жинайды. Бұл Pull моделінің нақты мысалы. Graphite бөлігінде керісінше жіберу "
                "сценарийі көрсетілді. Екі контрасты бір слайдта ұстау тақырыптың негізгі тұжырымын нығайтады."
            ),
            "screens": "Prometheus → Targets (UP) немесе /metrics үзіндісі (Fit AI стенді).",
        },
        {
            "n": 7,
            "title": "Zabbix стенді (Windows, Docker Compose)",
            "speakers": "Спикер 1",
            "slide_text": [
                "Каталог: monitoring/zabbix-lab",
                "Құрамы: MySQL + Zabbix Server + Web + Zabbix Agent",
                "Веб: http://localhost:8090",
                "Мақсат: кәсіби UI және агент арқылы дерек жинау",
            ],
            "script": (
                "Менің практикалық үлесім — Zabbix. Docker Compose арқылы көтерілді: MySQL, Zabbix Server, веб-интерфейс, "
                "Zabbix Agent. Windows Docker Desktop ортасында іске қосуға болады. Zabbixты тек push немесе тек pull деп "
                "қарапайым санатқа салу қиын — бұл гибридті платформа: агенттік деректер мен жоспарланған тексерулер бірге. "
                "Интерфейсте хост күйі және деректер ағыны расталады."
            ),
            "screens": "1) docker compose ps. 2) Zabbix — Dashboard немесе Hosts.",
        },
        {
            "n": 8,
            "title": "Zabbix: Docker нюансы және шешім",
            "speakers": "Спикер 1",
            "slide_text": [
                "Әдепкі: хост «Zabbix server» агентті 127.0.0.1 деп күтеді",
                "Docker: агент бөлек контейнерде → localhost жұмыс істемейді",
                "Шешім: Agent → Connect to: DNS, DNS name: zabbix-agent",
                "Нәтиже: Availability жасыл, Latest data толық",
            ],
            "script": (
                "Типтік мәселе: веб-интерфейсте әдепкі хост агентті сервердің ішіндегі localhost деп күтеді, ал Docker-да "
                "сервер мен агент әртүрлі контейнерде. Сондықтан agent is not available қатесі шықты. Шешім: Agent "
                "интерфейсінде DNS режимінде агент контейнерінің атауы көрсетілді. Содан кейін Availability жасыл, "
                "Latest data метрикалары жинала бастады. Бұл Docker мониторингінің практикалық ережесін көрсетеді."
            ),
            "screens": "1) Hosts — жасыл ZBX, zabbix-agent:10050. 2) Latest data — метрикалар.",
        },
        {
            "n": 9,
            "title": "Салыстыру кестесі",
            "speakers": "Спикер 2 (бірінші жарты) + Спикер 1 (екінші жарты)",
            "slide_text": [
                "Graphite — Push → Carbon — график Web-те",
                "Prometheus — Pull scrape — /metrics",
                "Nagios — белсенді HTTP тексеру — OK/200",
                "Zabbix — гибридті платформа — хост + Latest data",
            ],
            "script_sp2": (
                "Graphite бөлігінде Push идеясын жіберу арқылы көрсеттік. Nagios бөлігінде белсенді HTTP тексеруді көрсеттік."
            ),
            "script_sp1": (
                "Prometheus бойынша Pull scrape логикасын Fit AI жобасына сілтеме жасай отырып түсіндірдік. Zabbixта "
                "кәсіби интерфейс пен агенттік жинауды көрсеттік және Docker нюансын шештік."
            ),
            "screens": "Кесте толтырылған слайд; опционал — әр жүйенің кішкене скрині.",
        },
        {
            "n": 10,
            "title": "Қиындықтар және шешімдер",
            "speakers": "Спикер 1 + Спикер 2 (қысқа)",
            "slide_text": [
                "Порттар қақтығысы — әр стекте әртүрлі порттар",
                "Docker DNS — сервис атауымен қосылу",
                "Контейнерлерді бір-біріне дұрыс бағыттау",
            ],
            "script": (
                "Спикер 1: Менің жағымда негізгі қиындық Dockerдағы желілік логика болды: localhost емес, DNS арқылы "
                "агент атауы. Спикер 2: Серіктес жағыnda контейнерлерді дұрыс мекенжаймен байланыстыру және HTTP "
                "тексеруді дұрыс мақсатқа бағыттау маңызды болды."
            ),
            "screens": "Міндетті емес.",
        },
        {
            "n": 11,
            "title": "Қорытынды",
            "speakers": "Спикер 1 + Спикер 2 (бірге)",
            "slide_text": [
                "Pull және Push — әртүрлі сценарийлер үшін әртүрлі",
                "Практика теорияны растады",
                "Универсалды бір ғана «дұрыс» шешім жоқ",
            ],
            "script": (
                "Қорытындылай келе, біз мониторинг метрикаларын жинаудың әртүрлі механикасын нақты құралдармен көрсеттік. "
                "Универсалды бір ғана дұрыс шешім жоқ: маңыздысы — инфрақұрылымның түрі мен деректердің сипаты. "
                "Жұмысымыз тақырыптың практикалық сипатын растайды. Рахмет!"
            ),
            "screens": "Коллаж: Graphite + Nagios + Zabbix + Prometheus (кіші төрт бөлік).",
        },
        {
            "n": 12,
            "title": "Сұрақтар",
            "speakers": "Спикер 1 + Спикер 2",
            "slide_text": [
                "Спикер 1: Zabbix және Prometheus сұрақтарына",
                "Спикер 2: Graphite және Nagios сұрақтарына",
            ],
            "script": (
                "Сұрақтарыңызды қабылдаймыз. Мен Zabbix пен Prometheus жағына жауап беремін. Серіктесім Graphite пен "
                "Nagios жағына жауап береді."
            ),
            "screens": "Жоқ.",
        },
    ]

    add_heading(doc, "Презентация: Pull/Push — мәтін слайдтарға + оқу сценарийі (қазақша)", 0)

    add_para(
        doc,
        "Нұсқаулар: Әр төмендегі бөлім = бір слайд (Word-та бір бетке сыймаса, бетті өзіңіз бөлесіз). "
        "«Кім сөйлейді» жолы әр слайдтың басында көрсетілген.",
        bold=False,
    )
    doc.add_paragraph()

    for s in slides:
        add_heading(doc, f"Слайд {s['n']}: {s['title']}", level=1)
        add_para(doc, f"Кім сөйлейді: {s['speakers']}", bold=True)
        add_para(doc, "Слайдқа қоюға тезистер:", bold=True)
        for line in s["slide_text"]:
            add_para(doc, line)
        add_para(doc, "Оқуға сценарий (айту):", bold=True)
        if s["n"] == 9:
            add_para(doc, "Спикер 2: " + s["script_sp2"])
            add_para(doc, "Спикер 1: " + s["script_sp1"])
        else:
            add_para(doc, s["script"])
        add_para(doc, "Скриншоттар:", bold=True)
        add_para(doc, s["screens"])
        page_break(doc)

    # Кесте: кім қай слайдты айтады
    doc.add_paragraph()
    add_heading(doc, "Қосымша: спикерлер бойынша слайд нөмірлері", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Спикер"
    hdr[1].text = "Слайд нөмірлері"
    rows_data = [
        ("Спикер 1", "1 (бірге), 2, 6, 7, 8, 9 (екінші жарты), 10 (жарты), 11 (бірге), 12"),
        ("Спикер 2", "1 (бірге), 3, 4, 5, 9 (бірінші жарты), 10 (жарты), 11 (бірге), 12"),
    ]
    for a, b in rows_data:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b

    doc.save(OUT)
    print("OK:", OUT.encode("ascii", "replace").decode("ascii"))


if __name__ == "__main__":
    main()
